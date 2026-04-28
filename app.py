import streamlit as st
import pandas as pd
import numpy as np
import plotly.graph_objects as go
from scipy.signal import savgol_filter
from datetime import timedelta
import os
import zipfile
import io
from docx import Document
from docx.shared import Inches

st.set_page_config(layout="wide")

# SAFE KALEIDO CHECK
try:
    import kaleido
    KALEIDO_AVAILABLE = True
except:
    KALEIDO_AVAILABLE = False

# TITLE
st.title("Power Curve Analytics Report")

# LOGO
logo_path = os.path.join(os.path.dirname(__file__), "Envision.png")
col1, col2, col3 = st.columns([1,2,1])
with col2:
    if os.path.exists(logo_path):
        st.image(logo_path, width=300)

# SITE CAPACITY
SITE_CAPACITY = {site:3.3 for site in [
"CIP Hatalageri","JSW Tuljapur","Blupine Sagapara","Kalavad GJ","Kalavad_PH2",
"AMP_Energy","Wanki","CleanMax Motadevaliya","Ayana Amerli","Mahadev PH1",
"Blupine-I, Ambada-GJ","ACME Shapar","FP_Kudligi","Sprng TN",
"Otha Pithalpur-GJ","AMGEPL,Kurnool AP","ReNew1_Gadag","partner Ottapidaum",
"Cleanmax SANATHALI","Cleanmax Babra","RenfraEnergy Trichy","RENEW-03 Sholapur",
"Renew2 Chandwad","ReNew-4 Patoda","Clean max Jagalur","Sembcorp Tuticorin",
"Renew-4 Kudligi","Renew Otha","Cleanmax Honavad","Blueleaf Agar",
"JSW_Sandur","India_Hero_Doni"
]}

REF_FILE = "India site Standard & Theoretical PC data 1234.xlsx"
BIN_SIZE = 0.5
RATED_POWER = 3400.0

# SIDEBAR
st.sidebar.subheader("Upload SCADA File")
uploaded_file = st.sidebar.file_uploader("Upload SCADA CSV", type=["csv"])

if uploaded_file is None:
    st.warning("Please upload SCADA file")
    st.stop()

site = st.sidebar.selectbox("Select Site", list(SITE_CAPACITY.keys()))

mode = st.sidebar.radio(
    "Select View",
    ["Single Turbine", "Compare Turbines", "Show All Turbines"]
)

# LOAD SCADA
@st.cache_data
def load_scada(file):
    df = pd.read_csv(file, low_memory=False)
    df.columns = df.columns.str.strip()

    wind_col = [c for c in df.columns if "wind" in c.lower()][0]
    power_col = [c for c in df.columns if "power" in c.lower() or "active" in c.lower()][0]
    time_col = [c for c in df.columns if "time" in c.lower()][0]

    df[time_col] = pd.to_datetime(df[time_col], errors="coerce")
    df[wind_col] = pd.to_numeric(df[wind_col], errors="coerce")
    df[power_col] = pd.to_numeric(df[power_col], errors="coerce")

    df = df.dropna(subset=[wind_col,power_col,time_col])
    df["Name"] = df["Name"].astype(str).str.strip()

    return df, wind_col, power_col, time_col

df, wind_col, power_col, time_col = load_scada(uploaded_file)

# DATE FILTER
st.sidebar.markdown("Select Date Range")

max_date = df[time_col].max()

start_date = st.sidebar.date_input("Start Date", value=max_date - timedelta(days=15))
end_date = st.sidebar.date_input("End Date", value=max_date)

start_date = pd.to_datetime(start_date)
end_date = pd.to_datetime(end_date) + pd.Timedelta(days=1)

df = df[(df[time_col] >= start_date) & (df[time_col] <= end_date)]

# HEADER
num_turbines = df["Name"].nunique()
capacity_per_turbine = SITE_CAPACITY.get(site, 3.3)
total_capacity = num_turbines * capacity_per_turbine

st.subheader(f"{site} | {num_turbines} Turbines | {capacity_per_turbine} MW Each | Total: {round(total_capacity,2)} MW")
st.markdown(f"📅 Date Range: {start_date.date()} → {end_date.date()}")

# LOAD REFERENCE
@st.cache_data
def load_reference(site):
    ref_raw = pd.read_excel(REF_FILE, header=None)

    for r in range(ref_raw.shape[0]):
        for c in range(ref_raw.shape[1]):
            cell = str(ref_raw.iloc[r,c])
            if site.lower() in cell.lower():
                ref = ref_raw.iloc[r+2:r+60,[c-1,c+3]].copy()
                ref.columns=["WindSpeed","RefPower"]
                ref = ref.dropna()

                ref["WindSpeed"]=pd.to_numeric(ref["WindSpeed"], errors="coerce")
                ref["RefPower"]=pd.to_numeric(ref["RefPower"], errors="coerce")

                wind_bins = np.arange(4,10,BIN_SIZE)
                ref_interp = np.interp(wind_bins, ref["WindSpeed"], ref["RefPower"])

                return pd.DataFrame({"WindBin":wind_bins,"RefPower":ref_interp})

    st.error("Site not found")
    st.stop()

ref_curve = load_reference(site)

# PROCESS
def process_turbine(t):
    df_t = df[df["Name"]==t].copy()
    df_t = df_t[(df_t[wind_col]>=3)&(df_t[wind_col]<=25)&(df_t[power_col]>0)]

    if len(df_t)<30:
        return None

    std_dev = df_t[power_col].std()

    df_t["WindBin"] = (df_t[wind_col]/BIN_SIZE).round()*BIN_SIZE
    actual = df_t.groupby("WindBin").agg(AvgPower=(power_col,"mean")).reset_index()

    merged = ref_curve.merge(actual,on="WindBin",how="left")

    valid = merged["AvgPower"].notna()
    if valid.sum()>7:
        merged.loc[valid,"AvgPower"] = savgol_filter(merged.loc[valid,"AvgPower"],7,2)

    merged["Deviation_%"] = ((merged["AvgPower"]-merged["RefPower"])/merged["RefPower"])*100
    avg_dev = merged["Deviation_%"].mean(skipna=True)

    return df_t, merged, avg_dev, std_dev

# GRAPH
def plot_graph(df_t, merged, title, dev):
    fig = go.Figure()
    fig.add_trace(go.Scatter(x=df_t[wind_col],y=df_t[power_col],mode='markers'))
    fig.add_trace(go.Scatter(x=merged["WindBin"],y=merged["AvgPower"],mode='lines+markers'))
    fig.add_trace(go.Scatter(x=merged["WindBin"],y=merged["RefPower"],mode='lines',line=dict(dash='dash')))
    fig.update_layout(title=f"{title} (Dev: {round(dev,2)}%)")
    return fig

# COMMENT
def generate_comment(dev):
    if dev < -10:
        return "High Under"
    elif dev < -2:
        return "Under"
    elif dev > 8:
        return "High Over"
    elif dev > 2:
        return "Slight Over"
    else:
        return "Normal"

# MODE
turbines = df["Name"].unique()

if mode == "Single Turbine":
    turbines_to_show = [st.sidebar.selectbox("Select Turbine", turbines)]
elif mode == "Compare Turbines":
    turbines_to_show = st.sidebar.multiselect("Select Turbines", turbines)
else:
    turbines_to_show = turbines

# DISPLAY
cols = st.columns(2)
i = 0
results = []
zip_buffer = io.BytesIO()
zip_file = zipfile.ZipFile(zip_buffer, "w")

for t in turbines_to_show:
    res = process_turbine(t)
    if not res:
        continue

    df_t, merged, dev, std = res
    fig = plot_graph(df_t, merged, t, dev)

    with cols[i % 2]:
        st.plotly_chart(fig, use_container_width=True)
        st.code(generate_comment(dev))

    # save graph
    if KALEIDO_AVAILABLE:
        try:
            zip_file.writestr(f"{t}.png", fig.to_image(format="png"))
        except:
            pass

    results.append({
        "Turbine": t,
        "Deviation_%": round(dev,2),
        "Status": generate_comment(dev)
    })

    i += 1

# TABLE
st.subheader("Turbine Ranking")
results_df = pd.DataFrame(results).sort_values(by="Deviation_%")
st.dataframe(results_df, use_container_width=True)

# SAVE CSV
zip_file.writestr("report.csv", results_df.to_csv(index=False))

# ---------------- WORD REPORT ----------------
doc = Document()
doc.add_heading('Power Curve Analytics Report', 0)

doc.add_paragraph(f"Site: {site}")
doc.add_paragraph(f"Turbines: {num_turbines}")
doc.add_paragraph(f"Date: {start_date.date()} → {end_date.date()}")

if os.path.exists(logo_path):
    doc.add_picture(logo_path, width=Inches(2))

doc.add_heading("Turbine Analysis", 1)

for t in turbines_to_show:
    res = process_turbine(t)
    if not res:
        continue

    df_t, merged, dev, std = res
    doc.add_heading(t, 2)

    if KALEIDO_AVAILABLE:
        try:
            img_bytes = plot_graph(df_t, merged, t, dev).to_image(format="png")
            img_path = f"/tmp/{t}.png"
            with open(img_path, "wb") as f:
                f.write(img_bytes)
            doc.add_picture(img_path, width=Inches(5))
        except:
            pass

    doc.add_paragraph(generate_comment(dev))

doc.add_heading("Turbine Ranking", 1)

table = doc.add_table(rows=1, cols=3)
table.rows[0].cells[0].text = "Turbine"
table.rows[0].cells[1].text = "Deviation %"
table.rows[0].cells[2].text = "Status"

for _, r in results_df.iterrows():
    row_cells = table.add_row().cells
    row_cells[0].text = str(r["Turbine"])
    row_cells[1].text = str(r["Deviation_%"])
    row_cells[2].text = str(r["Status"])

doc_buffer = io.BytesIO()
doc.save(doc_buffer)

zip_file.writestr("PowerCurve_Report.docx", doc_buffer.getvalue())

zip_file.close()

# DOWNLOAD
st.download_button(
    "Download Dashboard (ZIP)",
    data=zip_buffer.getvalue(),
    file_name="PowerCurve_Full_Report.zip"
)
